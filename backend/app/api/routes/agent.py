# # import asyncio
# # import json
# # import queue
# # import uuid
# # from contextlib import suppress
# # from datetime import datetime
# # from typing import List

# # from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
# # from fastapi.responses import StreamingResponse
# # from sqlalchemy.orm import Session

# # from app.agents.graph_workflow import agent_workflow, build_initial_state
# # from app.core.file_utils import extract_file_payload
# # from app.core.json_utils import _safe_json_loads
# # from app.core.stream_logger import emit_done, emit_error, emit_final, emit_log, register_stream, unregister_stream
# # from app.db.database import get_db
# # from app.db.models import Conversation
# # from app.schemas.agent import AgentRunResponse
# # from app.kg.graph_store import kg_store

# # router = APIRouter(prefix="/agent", tags=["agent"])

# # def _payload_to_prompt_text(payload: dict) -> str:
# #     return f"\n\n--- 附件：{payload['filename']} 开始 ---\n{payload['text']}\n--- 附件：{payload['filename']} 结束 ---\n"

# # def _merge_analysis_snapshot(conversation: Conversation, payload: dict, current_mode: str) -> dict:
# #     snapshot = _safe_json_loads(conversation.analysis_snapshot, {})
# #     if current_mode not in snapshot:
# #         snapshot[current_mode] = {}
# #     snapshot[current_mode].update(payload)
# #     return snapshot

# # def _build_agent_response(
# #     final_state: dict,
# #     current_mode: str,
# #     thread_id: str,
# #     conversation: Conversation | None = None,
# # ) -> dict:
# #     analysis_snapshot = {}
# #     conversation_id = None
# #     document_status = "none"
# #     bound_file_name = None
# #     bound_file_uploaded_at = None
# #     if conversation is not None:
# #         conversation_id = conversation.id
# #         document_status = conversation.document_status or "none"
# #         bound_file_name = conversation.bound_file_name
# #         bound_file_uploaded_at = (
# #             conversation.bound_file_uploaded_at.isoformat()
# #             if conversation.bound_file_uploaded_at
# #             else None
# #         )
# #         analysis_snapshot = _safe_json_loads(conversation.analysis_snapshot, {})
# #     return {
# #         "current_mode": final_state.get("current_mode", current_mode),
# #         "selected_role": final_state.get("selected_role", ""),
# #         "critic_diagnosis": final_state.get("critic_diagnosis", {}),
# #         "planned_tasks": final_state.get("planned_tasks", []),
# #         "generated_content": final_state.get("generated_content", {}),
# #         "validation_status": final_state.get("validation_status", False),
# #         "validation_errors": final_state.get("validation_errors", ""),
# #         "attempt_count": final_state.get("attempt_count", 0),
# #         "thread_id": thread_id,
# #         "conversation_id": conversation_id,
# #         "document_status": document_status,
# #         "bound_file_name": bound_file_name,
# #         "bound_file_uploaded_at": bound_file_uploaded_at,
# #         "analysis_snapshot": analysis_snapshot,
# #         "hypergraph_data": final_state.get("hypergraph_data", {}), # 【新增】加入超图数据
# #         "stage_flow": final_state.get("stage_flow", {}),
# #         "kg_context": final_state.get("kg_context", {}),
# #     }

# # async def _load_file_payloads(files: List[UploadFile]) -> List[dict]:
# #     payloads = []
# #     for file in files:
# #         if file.filename:
# #             payloads.append(await extract_file_payload(file))
# #     return payloads

# # def _persist_bound_document(conversation: Conversation, payloads: List[dict], db: Session) -> Conversation:
# #     usable_payloads = [p for p in payloads if p.get("supported") and (p.get("text") or "").strip()]
# #     if not usable_payloads:
# #         return conversation
# #     merged_text = "\n\n".join(
# #         f"【附件：{p['filename']}】\n{p['text']}"
# #         for p in usable_payloads
# #     )
# #     conversation.bound_file_name = usable_payloads[0]["filename"]
# #     conversation.bound_document_text = merged_text
# #     conversation.bound_file_uploaded_at = datetime.utcnow()
# #     conversation.document_status = "bound"
# #     db.commit()
# #     db.refresh(conversation)
# #     return conversation

# # @router.get("/health")
# # def agent_health() -> dict:
# #     return {"status": "ok", "message": "agent route is running"}

# # @router.post("/run", response_model=AgentRunResponse)
# # async def run_agent(
# #     user_input: str = Form(default=""),
# #     current_mode: str = Form(default="learning"),
# #     thread_id: str = Form(default=""),
# #     conversation_id: str = Form(default=""),
# #     max_retry: int = Form(default=2),
# #     files: List[UploadFile] = File(default=[]),
# #     db: Session = Depends(get_db),
# # ) -> AgentRunResponse:
# #     try:
# #         if not thread_id:
# #             thread_id = str(uuid.uuid4())
# #         conversation = None
# #         if conversation_id:
# #             conversation = db.query(Conversation).filter(Conversation.id == conversation_id).first()
# #             if not conversation:
# #                 raise HTTPException(status_code=404, detail="会话不存在")
# #         file_payloads = await _load_file_payloads(files)

# #         if conversation is not None and file_payloads:
# #             conversation = _persist_bound_document(conversation, file_payloads, db)

# #         if conversation is not None and current_mode in {"project", "competition"}:
# #             if conversation.document_status != "bound":
# #                 raise HTTPException(
# #                     status_code=400,
# #                     detail="当前会话尚未绑定商业计划书，请先上传并绑定 BP 文档后再使用项目模式或竞赛模式。",
# #                 )

# #         file_texts = ""
# #         if conversation is None and file_payloads:
# #             file_texts = "".join(_payload_to_prompt_text(p) for p in file_payloads)
# #         combined_input = user_input
# #         if file_texts:
# #             combined_input = f"{user_input}{file_texts}".strip()
# #         if not combined_input:
# #             raise HTTPException(status_code=400, detail="输入文本与附件不能同时为空")

# #         initial_state = build_initial_state(
# #             user_input=combined_input,
# #             current_mode=current_mode,
# #             max_retry=max_retry,
# #             thread_id=thread_id,
# #             conversation_id=conversation.id if conversation else "",
# #             bound_document_text=conversation.bound_document_text if conversation else "",
# #             bound_file_name=conversation.bound_file_name if conversation else "",
# #             bound_file_uploaded_at=conversation.bound_file_uploaded_at.isoformat() if conversation and conversation.bound_file_uploaded_at else "",
# #             document_status=conversation.document_status if conversation else "none",
# #             analysis_snapshot=_safe_json_loads(conversation.analysis_snapshot, {}) if conversation else {},
# #         )
# #         config = {"configurable": {"thread_id": thread_id}}
# #         final_state = agent_workflow.invoke(initial_state, config=config)

# #         if conversation is not None:
# #             payload_before_snapshot = _build_agent_response(final_state, current_mode, thread_id, conversation)
# #             snapshot = _merge_analysis_snapshot(conversation, payload_before_snapshot, current_mode)
            
# #             # 【核心修改】确保提取到的 hypergraph_data / stage_flow 被持久化到数据库分析快照中
# #             if current_mode not in snapshot:
# #                 snapshot[current_mode] = {}
# #             if "hypergraph_data" in final_state and final_state["hypergraph_data"]:
# #                 snapshot[current_mode]["hypergraph_data"] = final_state["hypergraph_data"]
# #             if current_mode == "project" and "stage_flow" in final_state and final_state["stage_flow"]:
# #                 snapshot[current_mode]["stage_flow"] = final_state["stage_flow"]
# #             # ====== 新增：learning 模式保存 kg_context ======
# #             if current_mode == "learning" and "kg_context" in final_state and final_state["kg_context"]:
# #                 snapshot[current_mode]["kg_context"] = final_state["kg_context"]

# #             conversation.analysis_snapshot = json.dumps(snapshot, ensure_ascii=False)
# #             conversation.last_mode = current_mode
# #             db.commit()
# #             db.refresh(conversation)

# #         return AgentRunResponse(**_build_agent_response(final_state, current_mode, thread_id, conversation))
# #     except HTTPException:
# #         raise
# #     except Exception as e:
# #         raise HTTPException(status_code=500, detail=f"Agent workflow 执行失败：{str(e)}")

# # @router.post("/run/stream")
# # async def run_agent_stream(
# #     user_input: str = Form(default=""),
# #     current_mode: str = Form(default="learning"),
# #     thread_id: str = Form(default=""),
# #     conversation_id: str = Form(default=""),
# #     max_retry: int = Form(default=2),
# #     files: List[UploadFile] = File(default=[]),
# #     db: Session = Depends(get_db),
# # ):
# #     if not thread_id:
# #         thread_id = str(uuid.uuid4())
# #     conversation = None
# #     if conversation_id:
# #         conversation = db.query(Conversation).filter(Conversation.id == conversation_id).first()
# #         if not conversation:
# #             raise HTTPException(status_code=404, detail="会话不存在")
            
# #     file_payloads = await _load_file_payloads(files)
# #     if conversation is not None and file_payloads:
# #         conversation = _persist_bound_document(conversation, file_payloads, db)
        
# #     if conversation is not None and current_mode in {"project", "competition"}:
# #         if conversation.document_status != "bound":
# #             raise HTTPException(
# #                 status_code=400,
# #                 detail="当前会话尚未绑定商业计划书，请先上传并绑定 BP 文档后再使用项目模式或竞赛模式。",
# #             )
            
# #     file_texts = ""
# #     if conversation is None and file_payloads:
# #         file_texts = "".join(_payload_to_prompt_text(p) for p in file_payloads)
# #     combined_input = user_input
# #     if file_texts:
# #         combined_input = f"{user_input}{file_texts}".strip()
# #     if not combined_input:
# #         raise HTTPException(status_code=400, detail="输入文本与附件不能同时为空")

# #     initial_state = build_initial_state(
# #         user_input=combined_input,
# #         current_mode=current_mode,
# #         max_retry=max_retry,
# #         thread_id=thread_id,
# #         conversation_id=conversation.id if conversation else "",
# #         bound_document_text=conversation.bound_document_text if conversation else "",
# #         bound_file_name=conversation.bound_file_name if conversation else "",
# #         bound_file_uploaded_at=conversation.bound_file_uploaded_at.isoformat() if conversation and conversation.bound_file_uploaded_at else "",
# #         document_status=conversation.document_status if conversation else "none",
# #         analysis_snapshot=_safe_json_loads(conversation.analysis_snapshot, {}) if conversation else {},
# #     )
    
# #     config = {"configurable": {"thread_id": thread_id}}
# #     stream_queue = register_stream(thread_id)
    
# #     async def produce_workflow_events() -> None:
# #         try:
# #             emit_log(thread_id, "api", "请求已接收，开始创建工作流。")
# #             if conversation is not None and conversation.document_status == "bound":
# #                 emit_log(
# #                     thread_id,
# #                     "api",
# #                     f"当前会话已绑定文档：{conversation.bound_file_name or '未命名文件'}，将自动纳入本轮分析。",
# #                 )
# #             emit_log(thread_id, "api", "附件与输入已解析完成，准备启动多 Agent 协同流程。")
            
# #             final_state = await asyncio.to_thread(agent_workflow.invoke, initial_state, config=config)
            
# #             if conversation is not None:
# #                 payload_before_snapshot = _build_agent_response(final_state, current_mode, thread_id, conversation)
# #                 snapshot = _merge_analysis_snapshot(conversation, payload_before_snapshot, current_mode)
                
# #                 # 【核心修改】为流式输出同样保存超图信息与阶段推进状态
# #                 if current_mode not in snapshot:
# #                     snapshot[current_mode] = {}
# #                 if "hypergraph_data" in final_state and final_state["hypergraph_data"]:
# #                     snapshot[current_mode]["hypergraph_data"] = final_state["hypergraph_data"]
# #                 if current_mode == "project" and "stage_flow" in final_state and final_state["stage_flow"]:
# #                     snapshot[current_mode]["stage_flow"] = final_state["stage_flow"]
# #                 # ====== 新增：learning 模式保存 kg_context ======
# #                 if current_mode == "learning" and "kg_context" in final_state and final_state["kg_context"]:
# #                     snapshot[current_mode]["kg_context"] = final_state["kg_context"]
                    
# #                 conversation.analysis_snapshot = json.dumps(snapshot, ensure_ascii=False)
# #                 conversation.last_mode = current_mode
# #                 db.commit()
# #                 db.refresh(conversation)
                
# #             emit_log(thread_id, "api", "工作流执行完成，正在整理最终输出。")
# #             emit_final(thread_id, _build_agent_response(final_state, current_mode, thread_id, conversation))
# #         except Exception as e:
# #             emit_error(thread_id, f"Agent workflow 执行失败：{str(e)}")
# #         finally:
# #             emit_done(thread_id)

# #     producer_task = asyncio.create_task(produce_workflow_events())
    
# #     async def event_generator():
# #         try:
# #             while True:
# #                 try:
# #                     event = await asyncio.to_thread(stream_queue.get, True, 0.25)
# #                     yield f"data: {json.dumps(event, ensure_ascii=False, default=str)}\n\n"
# #                     if event.get("type") == "done":
# #                         break
# #                 except queue.Empty:
# #                     if producer_task.done():
# #                         await asyncio.sleep(0.05)
# #                     continue
# #         finally:
# #             with suppress(Exception):
# #                 await producer_task
# #             unregister_stream(thread_id)
            
# #     return StreamingResponse(
# #         event_generator(),
# #         media_type="text/event-stream",
# #         headers={
# #             "Cache-Control": "no-cache",
# #             "Connection": "keep-alive",
# #             "X-Accel-Buffering": "no",
# #         },
# #     )


# # @router.get("/kg/graph")
# # def get_learning_kg_graph(
# #     view: str = "query",
# #     q: str = "",
# # ) -> dict:
# #     """
# #     view=query  -> 当前查询子图
# #     view=global -> 全图
# #     q           -> learning 模式下的用户问题
# #     """
# #     try:
# #         if view == "global":
# #             return {
# #                 "ok": True,
# #                 "view": "global",
# #                 "data": kg_store.get_global_graph_for_visualization(limit=800),
# #             }

# #         all_db_nodes = kg_store.get_all_entity_names()
# #         detected_concepts = detect_concepts(q or "", all_db_nodes, fuzzy_threshold=72, debug=False)

# #         return {
# #             "ok": True,
# #             "view": "query",
# #             "data": kg_store.get_query_graph_for_visualization(detected_concepts),
# #         }
# #     except Exception as e:
# #         return {
# #             "ok": False,
# #             "view": view,
# #             "data": {"nodes": [], "triples": [], "hit_nodes": []},
# #             "error": str(e),
# #         }

# import asyncio
# import json
# import queue
# import uuid
# from contextlib import suppress
# from datetime import datetime
# from typing import List

# from fastapi import APIRouter, Body, Depends, File, Form, HTTPException, UploadFile
# from fastapi.responses import StreamingResponse
# from sqlalchemy.orm import Session

# from app.agents.graph_workflow import agent_workflow, build_initial_state
# from app.agents.mechanism.generator import generate_project_stage_draft_artifact
# from app.core.file_utils import extract_file_payload
# from app.core.json_utils import _safe_json_loads
# from app.core.stream_logger import emit_done, emit_error, emit_final, emit_log, register_stream, unregister_stream
# from app.db.database import get_db
# from app.db.models import Conversation
# from app.schemas.agent import AgentRunResponse, ProjectStageDraftResponse
# from app.kg.graph_store import kg_store
# from app.hypergraph.stage_config import get_stage_definition

# router = APIRouter(prefix="/agent", tags=["agent"])

# def _payload_to_prompt_text(payload: dict) -> str:
#     return f"\n\n--- 附件：{payload['filename']} 开始 ---\n{payload['text']}\n--- 附件：{payload['filename']} 结束 ---\n"

# def _merge_analysis_snapshot(conversation: Conversation, payload: dict, current_mode: str) -> dict:
#     snapshot = _safe_json_loads(conversation.analysis_snapshot, {})
#     if current_mode not in snapshot:
#         snapshot[current_mode] = {}
#     snapshot[current_mode].update(payload)
#     return snapshot


# def _build_project_chat_transcript(chat_history: list, target_stage_index: int) -> str:
#     if not chat_history:
#         return "无项目模式历史对话。"

#     transcript_items = []
#     for msg in chat_history:
#         if not isinstance(msg, dict):
#             continue

#         mode = msg.get("mode")
#         role = msg.get("role")
#         if mode not in {"project", "project_stage_draft"}:
#             continue

#         content = msg.get("content")
#         if not isinstance(content, dict):
#             content = {}

#         if content.get("stage_draft"):
#             continue

#         if role == "user":
#             text = msg.get("text")
#             text = str(text).strip() if text is not None else ""
#             if text:
#                 transcript_items.append(f"【学生】{text}")
#             continue

#         if role == "assistant":
#             reply = (content.get("reply") or msg.get("text") or "").strip()
#             if reply:
#                 transcript_items.append(f"【教练】{reply}")
#             flaw = (content.get("logic_flaw") or "").strip()
#             gap = (content.get("evidence_gap") or "").strip()
#             task = (content.get("only_one_task") or "").strip()
#             if flaw or gap or task:
#                 summary_bits = []
#                 if flaw:
#                     summary_bits.append(f"核心缺陷：{flaw}")
#                 if gap:
#                     summary_bits.append(f"证据缺口：{gap}")
#                 if task:
#                     summary_bits.append(f"本轮任务：{task}")
#                 transcript_items.append("【后台摘要】" + "；".join(summary_bits))

#     if not transcript_items:
#         return "无项目模式历史对话。"

#     # 只保留最近若干轮，避免提示词爆炸
#     transcript_items = transcript_items[-24:]
#     return "\n\n".join(transcript_items)



# def _build_agent_response(
#     final_state: dict,
#     current_mode: str,
#     thread_id: str,
#     conversation: Conversation | None = None,
# ) -> dict:
#     analysis_snapshot = {}
#     conversation_id = None
#     document_status = "none"
#     bound_file_name = None
#     bound_file_uploaded_at = None
#     if conversation is not None:
#         conversation_id = conversation.id
#         document_status = conversation.document_status or "none"
#         bound_file_name = conversation.bound_file_name
#         bound_file_uploaded_at = (
#             conversation.bound_file_uploaded_at.isoformat()
#             if conversation.bound_file_uploaded_at
#             else None
#         )
#         analysis_snapshot = _safe_json_loads(conversation.analysis_snapshot, {})
#     return {
#         "current_mode": final_state.get("current_mode", current_mode),
#         "selected_role": final_state.get("selected_role", ""),
#         "critic_diagnosis": final_state.get("critic_diagnosis", {}),
#         "planned_tasks": final_state.get("planned_tasks", []),
#         "generated_content": final_state.get("generated_content", {}),
#         "validation_status": final_state.get("validation_status", False),
#         "validation_errors": final_state.get("validation_errors", ""),
#         "attempt_count": final_state.get("attempt_count", 0),
#         "thread_id": thread_id,
#         "conversation_id": conversation_id,
#         "document_status": document_status,
#         "bound_file_name": bound_file_name,
#         "bound_file_uploaded_at": bound_file_uploaded_at,
#         "analysis_snapshot": analysis_snapshot,
#         "hypergraph_data": final_state.get("hypergraph_data", {}), # 【新增】加入超图数据
#         "stage_flow": final_state.get("stage_flow", {}),
#         "kg_context": final_state.get("kg_context", {}),
#     }

# async def _load_file_payloads(files: List[UploadFile]) -> List[dict]:
#     payloads = []
#     for file in files:
#         if file.filename:
#             payloads.append(await extract_file_payload(file))
#     return payloads

# def _persist_bound_document(conversation: Conversation, payloads: List[dict], db: Session) -> Conversation:
#     usable_payloads = [p for p in payloads if p.get("supported") and (p.get("text") or "").strip()]
#     if not usable_payloads:
#         return conversation
#     merged_text = "\n\n".join(
#         f"【附件：{p['filename']}】\n{p['text']}"
#         for p in usable_payloads
#     )
#     conversation.bound_file_name = usable_payloads[0]["filename"]
#     conversation.bound_document_text = merged_text
#     conversation.bound_file_uploaded_at = datetime.utcnow()
#     conversation.document_status = "bound"
#     db.commit()
#     db.refresh(conversation)
#     return conversation

# @router.get("/health")
# def agent_health() -> dict:
#     return {"status": "ok", "message": "agent route is running"}


# @router.post("/project-stage-draft", response_model=ProjectStageDraftResponse)
# def generate_project_stage_draft(
#     request: dict = Body(...),
#     db: Session = Depends(get_db),
# ) -> ProjectStageDraftResponse:
#     conversation_id = (request.get("conversation_id") or "").strip()
#     stage_id = (request.get("stage_id") or "").strip()

#     if not conversation_id:
#         raise HTTPException(status_code=400, detail="conversation_id 不能为空")
#     if not stage_id:
#         raise HTTPException(status_code=400, detail="stage_id 不能为空")

#     conversation = db.query(Conversation).filter(Conversation.id == conversation_id).first()
#     if not conversation:
#         raise HTTPException(status_code=404, detail="会话不存在")
#     if conversation.document_status != "bound" or not (conversation.bound_document_text or '').strip():
#         raise HTTPException(status_code=400, detail="当前会话尚未绑定商业计划书文档，无法生成阶段优化稿。")

#     snapshot = _safe_json_loads(conversation.analysis_snapshot, {})
#     project_snapshot = snapshot.get("project") or {}
#     stage_flow = project_snapshot.get("stage_flow") or {}
#     stages = stage_flow.get("stages") or {}
#     target_stage = stages.get(stage_id)

#     if not target_stage:
#         raise HTTPException(status_code=400, detail="当前会话尚未形成该阶段的项目推进快照，请先运行项目模式。")
#     if target_stage.get("status") != "passed":
#         raise HTTPException(status_code=400, detail="该阶段尚未通关，暂时不能生成阶段优化稿。")

#     stage_def = get_stage_definition(stage_id)
#     chat_history = _safe_json_loads(conversation.chat_history, [])
#     project_chat_transcript = _build_project_chat_transcript(chat_history, int(stage_def.get("index", 1)))
#     prior_artifact = ((project_snapshot.get("stage_artifacts") or {}).get(stage_id) or {})

#     try:
#         artifact = generate_project_stage_draft_artifact(
#             stage_id=stage_id,
#             stage_label=str(stage_def.get("label") or target_stage.get("label") or stage_id),
#             stage_goal=str(stage_def.get("goal") or target_stage.get("goal") or ''),
#             stage_rule_ids=list(stage_def.get("rule_ids") or []),
#             bound_document_text=conversation.bound_document_text or '',
#             project_chat_transcript=project_chat_transcript,
#             prior_artifact_text=str(prior_artifact.get("content") or ''),
#         )
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"阶段优化稿生成失败：{str(e)}")

#     artifact_payload = {
#         "stage_id": stage_id,
#         "stage_index": int(stage_def.get("index") or target_stage.get("index") or 0),
#         "stage_label": str(stage_def.get("label") or target_stage.get("label") or stage_id),
#         "title": artifact.get("title"),
#         "generation_notice": artifact.get("generation_notice"),
#         "content": artifact.get("content"),
#         "revision_highlights": artifact.get("revision_highlights") or [],
#         "generated_at": datetime.utcnow().isoformat(),
#     }

#     if "project" not in snapshot:
#         snapshot["project"] = {}
#     if "stage_artifacts" not in snapshot["project"] or not isinstance(snapshot["project"].get("stage_artifacts"), dict):
#         snapshot["project"]["stage_artifacts"] = {}
#     snapshot["project"]["stage_artifacts"][stage_id] = artifact_payload
#     conversation.analysis_snapshot = json.dumps(snapshot, ensure_ascii=False)
#     db.commit()
#     db.refresh(conversation)

#     return ProjectStageDraftResponse(
#         conversation_id=conversation.id,
#         stage_id=stage_id,
#         artifact=artifact_payload,
#         analysis_snapshot=_safe_json_loads(conversation.analysis_snapshot, {}),
#     )

# @router.post("/run", response_model=AgentRunResponse)
# async def run_agent(
#     user_input: str = Form(default=""),
#     current_mode: str = Form(default="learning"),
#     thread_id: str = Form(default=""),
#     conversation_id: str = Form(default=""),
#     max_retry: int = Form(default=2),
#     files: List[UploadFile] = File(default=[]),
#     db: Session = Depends(get_db),
# ) -> AgentRunResponse:
#     try:
#         if not thread_id:
#             thread_id = str(uuid.uuid4())
#         conversation = None
#         if conversation_id:
#             conversation = db.query(Conversation).filter(Conversation.id == conversation_id).first()
#             if not conversation:
#                 raise HTTPException(status_code=404, detail="会话不存在")
#         file_payloads = await _load_file_payloads(files)

#         if conversation is not None and file_payloads:
#             conversation = _persist_bound_document(conversation, file_payloads, db)

#         if conversation is not None and current_mode in {"project", "competition"}:
#             if conversation.document_status != "bound":
#                 raise HTTPException(
#                     status_code=400,
#                     detail="当前会话尚未绑定商业计划书，请先上传并绑定 BP 文档后再使用项目模式或竞赛模式。",
#                 )

#         file_texts = ""
#         if conversation is None and file_payloads:
#             file_texts = "".join(_payload_to_prompt_text(p) for p in file_payloads)
#         combined_input = user_input
#         if file_texts:
#             combined_input = f"{user_input}{file_texts}".strip()
#         if not combined_input:
#             raise HTTPException(status_code=400, detail="输入文本与附件不能同时为空")

#         initial_state = build_initial_state(
#             user_input=combined_input,
#             current_mode=current_mode,
#             max_retry=max_retry,
#             thread_id=thread_id,
#             conversation_id=conversation.id if conversation else "",
#             bound_document_text=conversation.bound_document_text if conversation else "",
#             bound_file_name=conversation.bound_file_name if conversation else "",
#             bound_file_uploaded_at=conversation.bound_file_uploaded_at.isoformat() if conversation and conversation.bound_file_uploaded_at else "",
#             document_status=conversation.document_status if conversation else "none",
#             analysis_snapshot=_safe_json_loads(conversation.analysis_snapshot, {}) if conversation else {},
#         )
#         config = {"configurable": {"thread_id": thread_id}}
#         final_state = agent_workflow.invoke(initial_state, config=config)

#         if conversation is not None:
#             payload_before_snapshot = _build_agent_response(final_state, current_mode, thread_id, conversation)
#             snapshot = _merge_analysis_snapshot(conversation, payload_before_snapshot, current_mode)
            
#             # 【核心修改】确保提取到的 hypergraph_data / stage_flow 被持久化到数据库分析快照中
#             if current_mode not in snapshot:
#                 snapshot[current_mode] = {}
#             if "hypergraph_data" in final_state and final_state["hypergraph_data"]:
#                 snapshot[current_mode]["hypergraph_data"] = final_state["hypergraph_data"]
#             if current_mode == "project" and "stage_flow" in final_state and final_state["stage_flow"]:
#                 snapshot[current_mode]["stage_flow"] = final_state["stage_flow"]
#             # ====== 新增：learning 模式保存 kg_context ======
#             if current_mode == "learning" and "kg_context" in final_state and final_state["kg_context"]:
#                 snapshot[current_mode]["kg_context"] = final_state["kg_context"]

#             conversation.analysis_snapshot = json.dumps(snapshot, ensure_ascii=False)
#             conversation.last_mode = current_mode
#             db.commit()
#             db.refresh(conversation)

#         return AgentRunResponse(**_build_agent_response(final_state, current_mode, thread_id, conversation))
#     except HTTPException:
#         raise
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Agent workflow 执行失败：{str(e)}")

# @router.post("/run/stream")
# async def run_agent_stream(
#     user_input: str = Form(default=""),
#     current_mode: str = Form(default="learning"),
#     thread_id: str = Form(default=""),
#     conversation_id: str = Form(default=""),
#     max_retry: int = Form(default=2),
#     files: List[UploadFile] = File(default=[]),
#     db: Session = Depends(get_db),
# ):
#     if not thread_id:
#         thread_id = str(uuid.uuid4())
#     conversation = None
#     if conversation_id:
#         conversation = db.query(Conversation).filter(Conversation.id == conversation_id).first()
#         if not conversation:
#             raise HTTPException(status_code=404, detail="会话不存在")
            
#     file_payloads = await _load_file_payloads(files)
#     if conversation is not None and file_payloads:
#         conversation = _persist_bound_document(conversation, file_payloads, db)
        
#     if conversation is not None and current_mode in {"project", "competition"}:
#         if conversation.document_status != "bound":
#             raise HTTPException(
#                 status_code=400,
#                 detail="当前会话尚未绑定商业计划书，请先上传并绑定 BP 文档后再使用项目模式或竞赛模式。",
#             )
            
#     file_texts = ""
#     if conversation is None and file_payloads:
#         file_texts = "".join(_payload_to_prompt_text(p) for p in file_payloads)
#     combined_input = user_input
#     if file_texts:
#         combined_input = f"{user_input}{file_texts}".strip()
#     if not combined_input:
#         raise HTTPException(status_code=400, detail="输入文本与附件不能同时为空")

#     initial_state = build_initial_state(
#         user_input=combined_input,
#         current_mode=current_mode,
#         max_retry=max_retry,
#         thread_id=thread_id,
#         conversation_id=conversation.id if conversation else "",
#         bound_document_text=conversation.bound_document_text if conversation else "",
#         bound_file_name=conversation.bound_file_name if conversation else "",
#         bound_file_uploaded_at=conversation.bound_file_uploaded_at.isoformat() if conversation and conversation.bound_file_uploaded_at else "",
#         document_status=conversation.document_status if conversation else "none",
#         analysis_snapshot=_safe_json_loads(conversation.analysis_snapshot, {}) if conversation else {},
#     )
    
#     config = {"configurable": {"thread_id": thread_id}}
#     stream_queue = register_stream(thread_id)
    
#     async def produce_workflow_events() -> None:
#         try:
#             emit_log(thread_id, "api", "请求已接收，开始创建工作流。")
#             if conversation is not None and conversation.document_status == "bound":
#                 emit_log(
#                     thread_id,
#                     "api",
#                     f"当前会话已绑定文档：{conversation.bound_file_name or '未命名文件'}，将自动纳入本轮分析。",
#                 )
#             emit_log(thread_id, "api", "附件与输入已解析完成，准备启动多 Agent 协同流程。")
            
#             final_state = await asyncio.to_thread(agent_workflow.invoke, initial_state, config=config)
            
#             if conversation is not None:
#                 payload_before_snapshot = _build_agent_response(final_state, current_mode, thread_id, conversation)
#                 snapshot = _merge_analysis_snapshot(conversation, payload_before_snapshot, current_mode)
                
#                 # 【核心修改】为流式输出同样保存超图信息与阶段推进状态
#                 if current_mode not in snapshot:
#                     snapshot[current_mode] = {}
#                 if "hypergraph_data" in final_state and final_state["hypergraph_data"]:
#                     snapshot[current_mode]["hypergraph_data"] = final_state["hypergraph_data"]
#                 if current_mode == "project" and "stage_flow" in final_state and final_state["stage_flow"]:
#                     snapshot[current_mode]["stage_flow"] = final_state["stage_flow"]
#                 # ====== 新增：learning 模式保存 kg_context ======
#                 if current_mode == "learning" and "kg_context" in final_state and final_state["kg_context"]:
#                     snapshot[current_mode]["kg_context"] = final_state["kg_context"]
                    
#                 conversation.analysis_snapshot = json.dumps(snapshot, ensure_ascii=False)
#                 conversation.last_mode = current_mode
#                 db.commit()
#                 db.refresh(conversation)
                
#             emit_log(thread_id, "api", "工作流执行完成，正在整理最终输出。")
#             emit_final(thread_id, _build_agent_response(final_state, current_mode, thread_id, conversation))
#         except Exception as e:
#             emit_error(thread_id, f"Agent workflow 执行失败：{str(e)}")
#         finally:
#             emit_done(thread_id)

#     producer_task = asyncio.create_task(produce_workflow_events())
    
#     async def event_generator():
#         try:
#             while True:
#                 try:
#                     event = await asyncio.to_thread(stream_queue.get, True, 0.25)
#                     yield f"data: {json.dumps(event, ensure_ascii=False, default=str)}\n\n"
#                     if event.get("type") == "done":
#                         break
#                 except queue.Empty:
#                     if producer_task.done():
#                         await asyncio.sleep(0.05)
#                     continue
#         finally:
#             with suppress(Exception):
#                 await producer_task
#             unregister_stream(thread_id)
            
#     return StreamingResponse(
#         event_generator(),
#         media_type="text/event-stream",
#         headers={
#             "Cache-Control": "no-cache",
#             "Connection": "keep-alive",
#             "X-Accel-Buffering": "no",
#         },
#     )


# @router.get("/kg/graph")
# def get_learning_kg_graph(
#     view: str = "query",
#     q: str = "",
# ) -> dict:
#     """
#     view=query  -> 当前查询子图
#     view=global -> 全图
#     q           -> learning 模式下的用户问题
#     """
#     try:
#         if view == "global":
#             return {
#                 "ok": True,
#                 "view": "global",
#                 "data": kg_store.get_global_graph_for_visualization(limit=800),
#             }

#         all_db_nodes = kg_store.get_all_entity_names()
#         detected_concepts = detect_concepts(q or "", all_db_nodes, fuzzy_threshold=72, debug=False)

#         return {
#             "ok": True,
#             "view": "query",
#             "data": kg_store.get_query_graph_for_visualization(detected_concepts),
#         }
#     except Exception as e:
#         return {
#             "ok": False,
#             "view": view,
#             "data": {"nodes": [], "triples": [], "hit_nodes": []},
#             "error": str(e),
#         }

import asyncio
import json
import queue
import re
import uuid
from contextlib import suppress
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from fastapi import APIRouter, Body, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from app.agents.graph_workflow import agent_workflow, build_initial_state
from app.agents.mechanism.generator import generate_final_project_book_artifact, generate_project_stage_draft_artifact
from app.core.file_utils import ATTACHMENT_ROOT, extract_file_payload
from app.core.json_utils import _safe_json_loads
from app.core.stream_logger import emit_done, emit_error, emit_event, emit_final, emit_log, register_stream, unregister_stream
from app.db.database import SessionLocal, get_db
from app.db.models import Conversation
from app.schemas.agent import AgentRunResponse, ProjectStageDraftResponse
from app.kg.graph_store import kg_store
from app.hypergraph.stage_config import get_stage_definition

router = APIRouter(prefix="/agent", tags=["agent"])

def _payload_to_prompt_text(payload: dict) -> str:
    return f"\n\n--- 附件：{payload['filename']} 开始 ---\n{payload['text']}\n--- 附件：{payload['filename']} 结束 ---\n"

def _merge_analysis_snapshot(conversation: Conversation, payload: dict, current_mode: str) -> dict:
    snapshot = _safe_json_loads(conversation.analysis_snapshot, {})
    if current_mode not in snapshot:
        snapshot[current_mode] = {}
    snapshot[current_mode].update(payload)
    return snapshot


def _build_project_chat_transcript(chat_history: list, target_stage_index: int) -> str:
    if not chat_history:
        return "无项目模式历史对话。"

    transcript_items = []
    for msg in chat_history:
        if not isinstance(msg, dict):
            continue

        mode = msg.get("mode")
        role = msg.get("role")
        if mode not in {"project", "project_stage_draft"}:
            continue

        content = msg.get("content")
        if not isinstance(content, dict):
            content = {}

        if content.get("stage_draft"):
            continue

        if role == "user":
            text = msg.get("text")
            text = str(text).strip() if text is not None else ""
            if text:
                transcript_items.append(f"【学生】{text}")
            continue

        if role == "assistant":
            reply = (content.get("reply") or msg.get("text") or "").strip()
            if reply:
                transcript_items.append(f"【教练】{reply}")
            flaw = (content.get("logic_flaw") or "").strip()
            gap = (content.get("evidence_gap") or "").strip()
            task = (content.get("only_one_task") or "").strip()
            if flaw or gap or task:
                summary_bits = []
                if flaw:
                    summary_bits.append(f"核心缺陷：{flaw}")
                if gap:
                    summary_bits.append(f"证据缺口：{gap}")
                if task:
                    summary_bits.append(f"本轮任务：{task}")
                transcript_items.append("【后台摘要】" + "；".join(summary_bits))

    if not transcript_items:
        return "无项目模式历史对话。"

    # 只保留最近若干轮，避免提示词爆炸
    transcript_items = transcript_items[-24:]
    return "\n\n".join(transcript_items)





def _build_all_project_stage_ids(stage_flow: dict) -> List[str]:
    stages = stage_flow.get("stages") or {}
    ordered = []
    for stage in stages.values():
        if not isinstance(stage, dict):
            continue
        if stage.get("status") != "passed":
            continue
        ordered.append((int(stage.get("index") or 0), str(stage.get("id") or "")))
    ordered = [item for item in ordered if item[1]]
    ordered.sort(key=lambda item: item[0])
    return [stage_id for _, stage_id in ordered]


def _is_valid_stage_artifact(artifact: dict | None) -> bool:
    if not isinstance(artifact, dict):
        return False
    return bool(str(artifact.get("content") or "").strip())


def _build_stage_artifact_payload(stage_def: dict, target_stage: dict, artifact: dict) -> dict:
    return {
        "stage_id": str(stage_def.get("id") or target_stage.get("id") or ''),
        "stage_index": int(stage_def.get("index") or target_stage.get("index") or 0),
        "stage_label": str(stage_def.get("label") or target_stage.get("label") or ''),
        "title": artifact.get("title"),
        "generation_notice": artifact.get("generation_notice"),
        "content": artifact.get("content"),
        "revision_highlights": artifact.get("revision_highlights") or [],
        "generated_at": datetime.utcnow().isoformat(),
    }


def _ensure_snapshot_project(snapshot: dict) -> dict:
    if "project" not in snapshot or not isinstance(snapshot.get("project"), dict):
        snapshot["project"] = {}
    if "stage_artifacts" not in snapshot["project"] or not isinstance(snapshot["project"].get("stage_artifacts"), dict):
        snapshot["project"]["stage_artifacts"] = {}
    return snapshot["project"]


def _emit_progress(thread_id: str, percent: int, message: str, step: str = '') -> None:
    emit_event(
        thread_id,
        {
            "type": "progress",
            "percent": max(0, min(int(percent), 100)),
            "message": message,
            "step": step or message,
        },
    )


def _sanitize_export_filename(name: str) -> str:
    safe = re.sub(r'[\/:*?"<>|]+', '_', str(name or '').strip())
    safe = re.sub(r'\s+', ' ', safe).strip().strip('.')
    return safe or '优化项目书'


def _append_docx_block(document: Document, text: str) -> None:
    stripped = (text or '').rstrip()
    if not stripped:
        document.add_paragraph('')
        return

    if stripped.startswith('### '):
        document.add_heading(stripped[4:].strip(), level=3)
        return
    if stripped.startswith('## '):
        document.add_heading(stripped[3:].strip(), level=2)
        return
    if stripped.startswith('# '):
        document.add_heading(stripped[2:].strip(), level=1)
        return

    bullet_match = re.match(r'^[-*•]\s+(.+)$', stripped)
    if bullet_match:
        document.add_paragraph(bullet_match.group(1).strip(), style='List Bullet')
        return

    number_match = re.match(r'^\d+[.)、]\s+(.+)$', stripped)
    if number_match:
        document.add_paragraph(number_match.group(1).strip(), style='List Number')
        return

    document.add_paragraph(stripped)


def _write_final_project_docx(*, title: str, generation_notice: str, content: str, revision_highlights: List[str], conversation_title: str, bound_file_name: str | None) -> dict:
    ATTACHMENT_ROOT.mkdir(parents=True, exist_ok=True)

    document = Document()
    try:
        normal_style = document.styles['Normal']
        normal_style.font.name = 'Microsoft YaHei'
        normal_style.font.size = Pt(10.5)
        if hasattr(normal_style, '_element') and hasattr(normal_style._element, 'rPr'):
            try:
                normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            except Exception:
                pass
    except Exception:
        pass

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title or '三阶段整合优化项目书（终稿整理版）')
    title_run.bold = True
    title_run.font.size = Pt(18)

    meta_paragraph = document.add_paragraph()
    meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_paragraph.add_run(f"导出时间：{datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC")
    meta_run.italic = True
    meta_run.font.size = Pt(9.5)

    if conversation_title or bound_file_name:
        intro = document.add_paragraph()
        intro.add_run('会话标题：').bold = True
        intro.add_run(conversation_title or '未命名会话')
        if bound_file_name:
            intro.add_run('    绑定文档：').bold = True
            intro.add_run(bound_file_name)

    if generation_notice:
        document.add_paragraph(generation_notice)

    if revision_highlights:
        document.add_heading('本次整合重点', level=1)
        for item in revision_highlights:
            cleaned = str(item or '').strip()
            if cleaned:
                document.add_paragraph(cleaned, style='List Bullet')

    document.add_heading('整合后项目书正稿', level=1)
    for block in re.split(r'\n\s*\n', content or ''):
        _append_docx_block(document, block)

    timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
    safe_stem = _sanitize_export_filename(title or conversation_title or '优化项目书')
    safe_name = f"export{timestamp}{uuid.uuid4().hex[:8]}_{safe_stem}.docx"
    file_path = ATTACHMENT_ROOT / safe_name
    document.save(str(file_path))

    return {
        "file_id": safe_name,
        "name": f"{safe_stem}.docx",
        "download_url": f"/api/v1/conversations/attachments/{safe_name}",
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "generated_at": datetime.utcnow().isoformat(),
        "path": str(file_path),
    }


def _generate_or_refresh_stage_artifact(snapshot: dict, conversation: Conversation, stage_id: str, *, project_chat_transcript: str) -> dict:
    project_snapshot = _ensure_snapshot_project(snapshot)
    stage_flow = project_snapshot.get("stage_flow") or {}
    stages = stage_flow.get("stages") or {}
    target_stage = stages.get(stage_id) or {}
    stage_def = get_stage_definition(stage_id)
    prior_artifact = (project_snapshot.get("stage_artifacts") or {}).get(stage_id) or {}

    artifact = generate_project_stage_draft_artifact(
        stage_id=stage_id,
        stage_label=str(stage_def.get("label") or target_stage.get("label") or stage_id),
        stage_goal=str(stage_def.get("goal") or target_stage.get("goal") or ''),
        stage_rule_ids=list(stage_def.get("rule_ids") or []),
        bound_document_text=conversation.bound_document_text or '',
        project_chat_transcript=project_chat_transcript,
        prior_artifact_text=str(prior_artifact.get("content") or ''),
    )
    artifact_payload = _build_stage_artifact_payload(stage_def, target_stage, artifact)
    project_snapshot["stage_artifacts"][stage_id] = artifact_payload
    return artifact_payload


def _run_project_final_export_job(conversation_id: str, thread_id: str) -> None:
    db = SessionLocal()
    try:
        _emit_progress(thread_id, 5, '正在校验会话与阶段状态…', 'validate')
        conversation = db.query(Conversation).filter(Conversation.id == conversation_id).first()
        if not conversation:
            raise ValueError('会话不存在')
        if conversation.document_status != 'bound' or not str(conversation.bound_document_text or '').strip():
            raise ValueError('当前会话尚未绑定商业计划书文档，无法导出最终优化项目书。')

        snapshot = _safe_json_loads(conversation.analysis_snapshot, {})
        project_snapshot = _ensure_snapshot_project(snapshot)
        stage_flow = project_snapshot.get('stage_flow') or {}
        stage_ids = _build_all_project_stage_ids(stage_flow)
        if not stage_ids:
            raise ValueError('当前会话尚未形成可导出的阶段成果，请先运行项目模式并完成阶段通关。')
        if stage_flow.get('overall_status') != 'completed' or len(stage_ids) < 3:
            raise ValueError('请先完成项目模式三阶段全部通关后，再导出最终优化项目书。')

        chat_history = _safe_json_loads(conversation.chat_history, [])
        project_chat_transcript = _build_project_chat_transcript(chat_history, 3)
        stage_artifacts = dict(project_snapshot.get('stage_artifacts') or {})

        progress_marks = [16, 32, 48]
        ordered_stage_artifacts: List[Dict[str, Any]] = []
        for idx, stage_id in enumerate(stage_ids[:3]):
            stage_def = get_stage_definition(stage_id)
            stage_label = str(stage_def.get('label') or stage_id)
            existing = stage_artifacts.get(stage_id)
            if _is_valid_stage_artifact(existing):
                _emit_progress(thread_id, progress_marks[idx], f'已读取{stage_label}增量优化稿。', f'stage_{idx + 1}_ready')
                artifact_payload = existing
            else:
                _emit_progress(thread_id, progress_marks[idx], f'正在补齐{stage_label}增量优化稿…', f'stage_{idx + 1}_generate')
                artifact_payload = _generate_or_refresh_stage_artifact(snapshot, conversation, stage_id, project_chat_transcript=project_chat_transcript)
                stage_artifacts[stage_id] = artifact_payload
                conversation.analysis_snapshot = json.dumps(snapshot, ensure_ascii=False)
                db.commit()
                db.refresh(conversation)
            ordered_stage_artifacts.append(artifact_payload)

        prior_export = (project_snapshot.get('final_export') or {}).get('content') or ''
        _emit_progress(thread_id, 72, '正在整合三阶段成果为最终正稿…', 'compose')
        final_artifact = generate_final_project_book_artifact(
            project_type_label=str(stage_flow.get('project_type_label') or '商业项目'),
            bound_document_text=conversation.bound_document_text or '',
            project_chat_transcript=project_chat_transcript,
            ordered_stage_artifacts=ordered_stage_artifacts,
            prior_export_text=str(prior_export or ''),
        )

        _emit_progress(thread_id, 88, '正在生成 Word 文档…', 'docx')
        file_payload = _write_final_project_docx(
            title=str(final_artifact.get('title') or '三阶段整合优化项目书（终稿整理版）'),
            generation_notice=str(final_artifact.get('generation_notice') or ''),
            content=str(final_artifact.get('content') or ''),
            revision_highlights=list(final_artifact.get('revision_highlights') or []),
            conversation_title=str(conversation.title or ''),
            bound_file_name=conversation.bound_file_name,
        )

        project_snapshot['final_export'] = {
            'title': final_artifact.get('title'),
            'generation_notice': final_artifact.get('generation_notice'),
            'content': final_artifact.get('content'),
            'revision_highlights': final_artifact.get('revision_highlights') or [],
            'generated_at': datetime.utcnow().isoformat(),
            'file': {k: v for k, v in file_payload.items() if k != 'path'},
        }
        conversation.analysis_snapshot = json.dumps(snapshot, ensure_ascii=False)
        db.commit()
        db.refresh(conversation)

        _emit_progress(thread_id, 100, '最终优化项目书已导出完成。', 'done')
        emit_final(
            thread_id,
            {
                'conversation_id': conversation.id,
                'analysis_snapshot': _safe_json_loads(conversation.analysis_snapshot, {}),
                'artifact': project_snapshot.get('final_export') or {},
                'file': {k: v for k, v in file_payload.items() if k != 'path'},
            },
        )
    except Exception as e:
        emit_error(thread_id, f'最终优化项目书导出失败：{str(e)}')
    finally:
        db.close()
        emit_done(thread_id)

def _build_agent_response(
    final_state: dict,
    current_mode: str,
    thread_id: str,
    conversation: Conversation | None = None,
) -> dict:
    analysis_snapshot = {}
    conversation_id = None
    document_status = "none"
    bound_file_name = None
    bound_file_uploaded_at = None
    if conversation is not None:
        conversation_id = conversation.id
        document_status = conversation.document_status or "none"
        bound_file_name = conversation.bound_file_name
        bound_file_uploaded_at = (
            conversation.bound_file_uploaded_at.isoformat()
            if conversation.bound_file_uploaded_at
            else None
        )
        analysis_snapshot = _safe_json_loads(conversation.analysis_snapshot, {})
    return {
        "current_mode": final_state.get("current_mode", current_mode),
        "selected_role": final_state.get("selected_role", ""),
        "critic_diagnosis": final_state.get("critic_diagnosis", {}),
        "planned_tasks": final_state.get("planned_tasks", []),
        "generated_content": final_state.get("generated_content", {}),
        "validation_status": final_state.get("validation_status", False),
        "validation_errors": final_state.get("validation_errors", ""),
        "attempt_count": final_state.get("attempt_count", 0),
        "thread_id": thread_id,
        "conversation_id": conversation_id,
        "document_status": document_status,
        "bound_file_name": bound_file_name,
        "bound_file_uploaded_at": bound_file_uploaded_at,
        "analysis_snapshot": analysis_snapshot,
        "hypergraph_data": final_state.get("hypergraph_data", {}), # 【新增】加入超图数据
        "stage_flow": final_state.get("stage_flow", {}),
        "kg_context": final_state.get("kg_context", {}),
    }

async def _load_file_payloads(files: List[UploadFile]) -> List[dict]:
    payloads = []
    for file in files:
        if file.filename:
            payloads.append(await extract_file_payload(file))
    return payloads

def _persist_bound_document(conversation: Conversation, payloads: List[dict], db: Session) -> Conversation:
    usable_payloads = [p for p in payloads if p.get("supported") and (p.get("text") or "").strip()]
    if not usable_payloads:
        return conversation
    merged_text = "\n\n".join(
        f"【附件：{p['filename']}】\n{p['text']}"
        for p in usable_payloads
    )
    conversation.bound_file_name = usable_payloads[0]["filename"]
    conversation.bound_document_text = merged_text
    conversation.bound_file_uploaded_at = datetime.utcnow()
    conversation.document_status = "bound"
    db.commit()
    db.refresh(conversation)
    return conversation

@router.get("/health")
def agent_health() -> dict:
    return {"status": "ok", "message": "agent route is running"}


@router.post("/project-stage-draft", response_model=ProjectStageDraftResponse)
def generate_project_stage_draft(
    request: dict = Body(...),
    db: Session = Depends(get_db),
) -> ProjectStageDraftResponse:
    conversation_id = (request.get("conversation_id") or "").strip()
    stage_id = (request.get("stage_id") or "").strip()

    if not conversation_id:
        raise HTTPException(status_code=400, detail="conversation_id 不能为空")
    if not stage_id:
        raise HTTPException(status_code=400, detail="stage_id 不能为空")

    conversation = db.query(Conversation).filter(Conversation.id == conversation_id).first()
    if not conversation:
        raise HTTPException(status_code=404, detail="会话不存在")
    if conversation.document_status != "bound" or not (conversation.bound_document_text or '').strip():
        raise HTTPException(status_code=400, detail="当前会话尚未绑定商业计划书文档，无法生成阶段优化稿。")

    snapshot = _safe_json_loads(conversation.analysis_snapshot, {})
    project_snapshot = snapshot.get("project") or {}
    stage_flow = project_snapshot.get("stage_flow") or {}
    stages = stage_flow.get("stages") or {}
    target_stage = stages.get(stage_id)

    if not target_stage:
        raise HTTPException(status_code=400, detail="当前会话尚未形成该阶段的项目推进快照，请先运行项目模式。")
    if target_stage.get("status") != "passed":
        raise HTTPException(status_code=400, detail="该阶段尚未通关，暂时不能生成阶段优化稿。")

    stage_def = get_stage_definition(stage_id)
    chat_history = _safe_json_loads(conversation.chat_history, [])
    project_chat_transcript = _build_project_chat_transcript(chat_history, int(stage_def.get("index", 1)))
    prior_artifact = ((project_snapshot.get("stage_artifacts") or {}).get(stage_id) or {})

    try:
        artifact = generate_project_stage_draft_artifact(
            stage_id=stage_id,
            stage_label=str(stage_def.get("label") or target_stage.get("label") or stage_id),
            stage_goal=str(stage_def.get("goal") or target_stage.get("goal") or ''),
            stage_rule_ids=list(stage_def.get("rule_ids") or []),
            bound_document_text=conversation.bound_document_text or '',
            project_chat_transcript=project_chat_transcript,
            prior_artifact_text=str(prior_artifact.get("content") or ''),
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"阶段优化稿生成失败：{str(e)}")

    artifact_payload = {
        "stage_id": stage_id,
        "stage_index": int(stage_def.get("index") or target_stage.get("index") or 0),
        "stage_label": str(stage_def.get("label") or target_stage.get("label") or stage_id),
        "title": artifact.get("title"),
        "generation_notice": artifact.get("generation_notice"),
        "content": artifact.get("content"),
        "revision_highlights": artifact.get("revision_highlights") or [],
        "generated_at": datetime.utcnow().isoformat(),
    }

    if "project" not in snapshot:
        snapshot["project"] = {}
    if "stage_artifacts" not in snapshot["project"] or not isinstance(snapshot["project"].get("stage_artifacts"), dict):
        snapshot["project"]["stage_artifacts"] = {}
    snapshot["project"]["stage_artifacts"][stage_id] = artifact_payload
    conversation.analysis_snapshot = json.dumps(snapshot, ensure_ascii=False)
    db.commit()
    db.refresh(conversation)

    return ProjectStageDraftResponse(
        conversation_id=conversation.id,
        stage_id=stage_id,
        artifact=artifact_payload,
        analysis_snapshot=_safe_json_loads(conversation.analysis_snapshot, {}),
    )



@router.post("/project-final-export/stream")
async def generate_project_final_export_stream(
    request: dict = Body(...),
):
    conversation_id = str((request or {}).get("conversation_id") or '').strip()
    if not conversation_id:
        raise HTTPException(status_code=400, detail="conversation_id 不能为空")

    thread_id = str((request or {}).get("thread_id") or '').strip() or str(uuid.uuid4())
    stream_queue = register_stream(thread_id)
    producer_task = asyncio.create_task(asyncio.to_thread(_run_project_final_export_job, conversation_id, thread_id))

    async def event_generator():
        try:
            while True:
                try:
                    event = await asyncio.to_thread(stream_queue.get, True, 0.25)
                    yield f"data: {json.dumps(event, ensure_ascii=False, default=str)}\n\n"
                    if event.get("type") == "done":
                        break
                except queue.Empty:
                    if producer_task.done():
                        await asyncio.sleep(0.05)
                    continue
        finally:
            with suppress(Exception):
                await producer_task
            unregister_stream(thread_id)

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )

@router.post("/run", response_model=AgentRunResponse)
async def run_agent(
    user_input: str = Form(default=""),
    current_mode: str = Form(default="learning"),
    thread_id: str = Form(default=""),
    conversation_id: str = Form(default=""),
    max_retry: int = Form(default=2),
    files: List[UploadFile] = File(default=[]),
    db: Session = Depends(get_db),
) -> AgentRunResponse:
    try:
        if not thread_id:
            thread_id = str(uuid.uuid4())
        conversation = None
        if conversation_id:
            conversation = db.query(Conversation).filter(Conversation.id == conversation_id).first()
            if not conversation:
                raise HTTPException(status_code=404, detail="会话不存在")
        file_payloads = await _load_file_payloads(files)

        if conversation is not None and file_payloads:
            conversation = _persist_bound_document(conversation, file_payloads, db)

        if conversation is not None and current_mode in {"project", "competition"}:
            if conversation.document_status != "bound":
                raise HTTPException(
                    status_code=400,
                    detail="当前会话尚未绑定商业计划书，请先上传并绑定 BP 文档后再使用项目模式或竞赛模式。",
                )

        file_texts = ""
        if conversation is None and file_payloads:
            file_texts = "".join(_payload_to_prompt_text(p) for p in file_payloads)
        combined_input = user_input
        if file_texts:
            combined_input = f"{user_input}{file_texts}".strip()
        if not combined_input:
            raise HTTPException(status_code=400, detail="输入文本与附件不能同时为空")

        initial_state = build_initial_state(
            user_input=combined_input,
            current_mode=current_mode,
            max_retry=max_retry,
            thread_id=thread_id,
            conversation_id=conversation.id if conversation else "",
            bound_document_text=conversation.bound_document_text if conversation else "",
            bound_file_name=conversation.bound_file_name if conversation else "",
            bound_file_uploaded_at=conversation.bound_file_uploaded_at.isoformat() if conversation and conversation.bound_file_uploaded_at else "",
            document_status=conversation.document_status if conversation else "none",
            analysis_snapshot=_safe_json_loads(conversation.analysis_snapshot, {}) if conversation else {},
        )
        config = {"configurable": {"thread_id": thread_id}}
        final_state = agent_workflow.invoke(initial_state, config=config)

        if conversation is not None:
            payload_before_snapshot = _build_agent_response(final_state, current_mode, thread_id, conversation)
            snapshot = _merge_analysis_snapshot(conversation, payload_before_snapshot, current_mode)
            
            # 【核心修改】确保提取到的 hypergraph_data / stage_flow 被持久化到数据库分析快照中
            if current_mode not in snapshot:
                snapshot[current_mode] = {}
            if "hypergraph_data" in final_state and final_state["hypergraph_data"]:
                snapshot[current_mode]["hypergraph_data"] = final_state["hypergraph_data"]
            if current_mode == "project" and "stage_flow" in final_state and final_state["stage_flow"]:
                snapshot[current_mode]["stage_flow"] = final_state["stage_flow"]
            # ====== 新增：learning 模式保存 kg_context ======
            if current_mode == "learning" and "kg_context" in final_state and final_state["kg_context"]:
                snapshot[current_mode]["kg_context"] = final_state["kg_context"]

            conversation.analysis_snapshot = json.dumps(snapshot, ensure_ascii=False)
            conversation.last_mode = current_mode
            db.commit()
            db.refresh(conversation)

        return AgentRunResponse(**_build_agent_response(final_state, current_mode, thread_id, conversation))
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Agent workflow 执行失败：{str(e)}")

@router.post("/run/stream")
async def run_agent_stream(
    user_input: str = Form(default=""),
    current_mode: str = Form(default="learning"),
    thread_id: str = Form(default=""),
    conversation_id: str = Form(default=""),
    max_retry: int = Form(default=2),
    files: List[UploadFile] = File(default=[]),
    db: Session = Depends(get_db),
):
    if not thread_id:
        thread_id = str(uuid.uuid4())
    conversation = None
    if conversation_id:
        conversation = db.query(Conversation).filter(Conversation.id == conversation_id).first()
        if not conversation:
            raise HTTPException(status_code=404, detail="会话不存在")
            
    file_payloads = await _load_file_payloads(files)
    if conversation is not None and file_payloads:
        conversation = _persist_bound_document(conversation, file_payloads, db)
        
    if conversation is not None and current_mode in {"project", "competition"}:
        if conversation.document_status != "bound":
            raise HTTPException(
                status_code=400,
                detail="当前会话尚未绑定商业计划书，请先上传并绑定 BP 文档后再使用项目模式或竞赛模式。",
            )
            
    file_texts = ""
    if conversation is None and file_payloads:
        file_texts = "".join(_payload_to_prompt_text(p) for p in file_payloads)
    combined_input = user_input
    if file_texts:
        combined_input = f"{user_input}{file_texts}".strip()
    if not combined_input:
        raise HTTPException(status_code=400, detail="输入文本与附件不能同时为空")

    initial_state = build_initial_state(
        user_input=combined_input,
        current_mode=current_mode,
        max_retry=max_retry,
        thread_id=thread_id,
        conversation_id=conversation.id if conversation else "",
        bound_document_text=conversation.bound_document_text if conversation else "",
        bound_file_name=conversation.bound_file_name if conversation else "",
        bound_file_uploaded_at=conversation.bound_file_uploaded_at.isoformat() if conversation and conversation.bound_file_uploaded_at else "",
        document_status=conversation.document_status if conversation else "none",
        analysis_snapshot=_safe_json_loads(conversation.analysis_snapshot, {}) if conversation else {},
    )
    
    config = {"configurable": {"thread_id": thread_id}}
    stream_queue = register_stream(thread_id)
    
    async def produce_workflow_events() -> None:
        try:
            emit_log(thread_id, "api", "请求已接收，开始创建工作流。")
            if conversation is not None and conversation.document_status == "bound":
                emit_log(
                    thread_id,
                    "api",
                    f"当前会话已绑定文档：{conversation.bound_file_name or '未命名文件'}，将自动纳入本轮分析。",
                )
            emit_log(thread_id, "api", "附件与输入已解析完成，准备启动多 Agent 协同流程。")
            
            final_state = await asyncio.to_thread(agent_workflow.invoke, initial_state, config=config)
            
            if conversation is not None:
                payload_before_snapshot = _build_agent_response(final_state, current_mode, thread_id, conversation)
                snapshot = _merge_analysis_snapshot(conversation, payload_before_snapshot, current_mode)
                
                # 【核心修改】为流式输出同样保存超图信息与阶段推进状态
                if current_mode not in snapshot:
                    snapshot[current_mode] = {}
                if "hypergraph_data" in final_state and final_state["hypergraph_data"]:
                    snapshot[current_mode]["hypergraph_data"] = final_state["hypergraph_data"]
                if current_mode == "project" and "stage_flow" in final_state and final_state["stage_flow"]:
                    snapshot[current_mode]["stage_flow"] = final_state["stage_flow"]
                # ====== 新增：learning 模式保存 kg_context ======
                if current_mode == "learning" and "kg_context" in final_state and final_state["kg_context"]:
                    snapshot[current_mode]["kg_context"] = final_state["kg_context"]
                    
                conversation.analysis_snapshot = json.dumps(snapshot, ensure_ascii=False)
                conversation.last_mode = current_mode
                db.commit()
                db.refresh(conversation)
                
            emit_log(thread_id, "api", "工作流执行完成，正在整理最终输出。")
            emit_final(thread_id, _build_agent_response(final_state, current_mode, thread_id, conversation))
        except Exception as e:
            emit_error(thread_id, f"Agent workflow 执行失败：{str(e)}")
        finally:
            emit_done(thread_id)

    producer_task = asyncio.create_task(produce_workflow_events())
    
    async def event_generator():
        try:
            while True:
                try:
                    event = await asyncio.to_thread(stream_queue.get, True, 0.25)
                    yield f"data: {json.dumps(event, ensure_ascii=False, default=str)}\n\n"
                    if event.get("type") == "done":
                        break
                except queue.Empty:
                    if producer_task.done():
                        await asyncio.sleep(0.05)
                    continue
        finally:
            with suppress(Exception):
                await producer_task
            unregister_stream(thread_id)
            
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@router.get("/kg/graph")
def get_learning_kg_graph(
    view: str = "query",
    q: str = "",
) -> dict:
    """
    view=query  -> 当前查询子图
    view=global -> 全图
    q           -> learning 模式下的用户问题
    """
    try:
        if view == "global":
            return {
                "ok": True,
                "view": "global",
                "data": kg_store.get_global_graph_for_visualization(limit=800),
            }

        all_db_nodes = kg_store.get_all_entity_names()
        detected_concepts = detect_concepts(q or "", all_db_nodes, fuzzy_threshold=72, debug=False)

        return {
            "ok": True,
            "view": "query",
            "data": kg_store.get_query_graph_for_visualization(detected_concepts),
        }
    except Exception as e:
        return {
            "ok": False,
            "view": view,
            "data": {"nodes": [], "triples": [], "hit_nodes": []},
            "error": str(e),
        }